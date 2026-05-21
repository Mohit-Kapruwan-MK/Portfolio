"""
Generates Mohit Kapruwan's resume as both PDF and DOCX.
Run: python3 generate_resume.py
"""

import os
import sys

RESUME_DIR = os.path.dirname(os.path.abspath(__file__))
PDF_OUT  = os.path.join(RESUME_DIR, "resume.pdf")
DOCX_OUT = os.path.join(RESUME_DIR, "resume.docx")

# ─────────────────────────────────────────────────────────────────────────────
#  DATA
# ─────────────────────────────────────────────────────────────────────────────

CONTACT = {
    "name":     "Mohit Kapruwan",
    "title":    "Flutter Developer  ·  Mobile & Cross-Platform",
    "email":    "mohitkapruwan4@gmail.com",
    "phone":    "+91 97199 72748",
    "location": "Mohali, India",
    "linkedin": "linkedin.com/in/mohit-kapruwan-96b62b22b",
    "github":   "github.com/Mohit-Kapruwan-MK",
}

SUMMARY = (
    "Passionate and results-driven Software Engineer with 3+ years of experience in "
    "mobile app development, specialising in Flutter. Skilled in building high-performance, "
    "scalable, and user-friendly cross-platform applications for Android and iOS using "
    "modern UI/UX principles and Clean Architecture. Delivered 4 production apps live on "
    "the Google Play Store and Apple App Store. Currently open to full-time or freelance "
    "opportunities to contribute, grow, and create meaningful user experiences."
)

EXPERIENCE = [
    {
        "role":    "Software Developer",
        "company": "Live Informatics Deftsoft PVT LTD",
        "location":"Mohali, India",
        "period":  "July 2023 – Present",
        "bullets": [
            "Develop and maintain cross-platform mobile applications for Android and iOS using Flutter and Dart; own the full release cycle through App Store and Play Store publication.",
            "Architect scalable solutions following Clean Architecture, SOLID principles, and TDD, producing maintainable and testable codebases.",
            "Integrate Firebase (Auth, Firestore, Cloud Messaging), REST APIs, and third-party SDKs including TT Lock, Google Maps, and payment gateways.",
            "Implement real-time features — live chat, push notifications, match algorithms, and booking flows — across multiple product verticals.",
            "Collaborate with cross-functional teams using Agile methodologies; deliver pixel-perfect UI from Figma designs.",
            "Optimise app performance, reduce build size, and resolve memory issues to ensure smooth, production-ready experiences.",
        ],
    },
    {
        "role":    "Customer Service Associate",
        "company": "iEnergizer (Pine Labs)",
        "location":"Noida, India",
        "period":  "June 2022 – December 2022",
        "bullets": [
            "Provided end-to-end customer and technical support for Pine Labs POS platforms, handling account queries and payment transaction issues for merchants across India.",
            "Developed strong communication and problem-solving skills by resolving escalated issues efficiently, maintaining high customer-satisfaction scores.",
        ],
    },
]

PROJECTS = [
    {
        "title":    "Dealsquawk",
        "store":    "Apple App Store",
        "link":     "apps.apple.com/us/app/dealsquawk/id6746202322",
        "desc":     "Social community app — users create nests, share posts, interact with communities, and chat in real-time.",
        "stack":    "Flutter · Firebase Auth · Cloud Firestore · Real-time Chat · Push Notifications · REST APIs · State Management",
    },
    {
        "title":    "YezidiLink",
        "store":    "Google Play",
        "link":     "play.google.com/store/apps/details?id=com.app.yezidilink",
        "desc":     "Dating app with swipe-based matching, profile discovery, real-time chat, and user connection features.",
        "stack":    "Flutter · Firebase · Real-time Database · Match Algorithm · Chat Integration · Push Notifications · REST APIs",
    },
    {
        "title":    "Modyaf",
        "store":    "Google Play",
        "link":     "play.google.com/store/apps/details?id=com.modyaf.host",
        "desc":     "Airbnb-style rental booking platform with smart room access via TT Lock SDK, Google Maps, and full reservation management.",
        "stack":    "Flutter · Firebase · REST APIs · Google Maps · Payment Gateway · TT Lock SDK · Push Notifications",
    },
    {
        "title":    "SuperNotes AI",
        "store":    "Google Play",
        "link":     "play.google.com/store/apps/details?id=com.supernotes.app",
        "desc":     "AI-powered note-taking app with audio-to-text transcription, offline local storage, and drag-and-drop note organisation.",
        "stack":    "Flutter · Local Database · Speech-to-Text · AI Processing · Offline Storage · Drag & Drop UI · Provider",
    },
]

SKILLS = [
    ("Flutter",             95),
    ("Dart",                90),
    ("REST APIs",           90),
    ("Provider / Riverpod", 88),
    ("Firebase",            85),
    ("Git",                 85),
    ("UI/UX Design",        80),
    ("CI/CD",               75),
]

TOOLS = [
    "Flutter", "Dart", "Firebase", "Riverpod", "Provider",
    "React", "Next.js", "Node.js", "Express", "MongoDB",
    "MySQL", "TypeScript", "SQLite", "REST APIs", "GraphQL", "Git",
]

PRACTICES = [
    "Clean Architecture", "SOLID Principles", "TDD",
    "Agile / Scrum", "State Management", "Performance Optimisation",
    "App Store Deployment", "Push Notifications",
]

LANGUAGES = [("English", "Professional"), ("Hindi", "Native")]


# ─────────────────────────────────────────────────────────────────────────────
#  1. PDF via WeasyPrint
# ─────────────────────────────────────────────────────────────────────────────

HTML_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

@page { size: A4; margin: 0; }

*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

body {
  font-family: 'Inter', Helvetica, Arial, sans-serif;
  background: #fff;
  color: #0F172A;
  width: 210mm;
}

:root {
  --accent: #0D9488;
  --accent-light: #CCFBF1;
  --muted: #64748B;
  --border: #E2E8F0;
  --sidebar-bg: #F8FAFC;
}

/* Header */
.header {
  padding: 26px 28px 22px 26px;
  border-bottom: 2.5px solid #0D9488;
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 20px;
}
.h-name { font-size: 28px; font-weight: 800; letter-spacing: -0.5px; }
.h-name span { color: #0D9488; }
.h-sub { font-size: 11px; font-weight: 500; color: #64748B; letter-spacing: 2px; text-transform: uppercase; margin-top: 4px; }
.h-contacts { text-align: right; }
.h-contacts div { font-size: 10.5px; color: #64748B; margin-bottom: 3px; }

/* Body */
.body { display: flex; }

/* Sidebar */
.sidebar {
  width: 67mm;
  background: #F8FAFC;
  border-right: 1px solid #E2E8F0;
  padding: 24px 18px 24px 20px;
}
.s-section { margin-bottom: 22px; }

.s-title {
  font-size: 8.5px;
  font-weight: 700;
  letter-spacing: 2.5px;
  text-transform: uppercase;
  color: #0D9488;
  border-bottom: 1px solid #E2E8F0;
  padding-bottom: 5px;
  margin-bottom: 10px;
}

/* Skill bars */
.skill { margin-bottom: 8px; }
.skill-row { display: flex; justify-content: space-between; margin-bottom: 3px; }
.skill-name { font-size: 10.5px; font-weight: 500; }
.skill-pct { font-size: 9.5px; color: #64748B; }
.track { height: 3.5px; background: #E2E8F0; border-radius: 2px; }
.fill { height: 100%; background: linear-gradient(to right, #0D9488, #14B8A6); border-radius: 2px; }

/* Tags */
.tags { display: flex; flex-wrap: wrap; gap: 4px; }
.tag { font-size: 9px; font-weight: 600; color: #0D9488; background: #CCFBF1; padding: 2px 6px; border-radius: 3px; }
.tag-g { font-size: 9px; font-weight: 500; color: #64748B; background: #F1F5F9; padding: 2px 6px; border-radius: 3px; }

/* Lang */
.lang-row { display: flex; justify-content: space-between; font-size: 10.5px; margin-bottom: 5px; }
.lang-name { font-weight: 600; }
.lang-lv { color: #64748B; }

/* Main */
.main { flex: 1; padding: 24px 26px 24px 22px; }
.m-section { margin-bottom: 20px; }
.m-title {
  font-size: 8.5px; font-weight: 700; letter-spacing: 2.5px;
  text-transform: uppercase; color: #0D9488;
  border-bottom: 1px solid #E2E8F0; padding-bottom: 5px; margin-bottom: 12px;
}

/* Summary */
.summary { font-size: 11.5px; color: #334155; line-height: 1.7; }

/* Experience */
.exp { margin-bottom: 14px; }
.exp-head { display: flex; justify-content: space-between; align-items: flex-start; gap: 8px; margin-bottom: 1px; }
.exp-role { font-size: 12px; font-weight: 700; }
.exp-period { font-size: 9.5px; font-weight: 600; color: #0D9488; background: #CCFBF1; padding: 2px 8px; border-radius: 20px; white-space: nowrap; }
.exp-co { font-size: 10.5px; color: #64748B; font-weight: 500; margin-bottom: 6px; }
.bullets { list-style: none; padding: 0; }
.bullets li { font-size: 10.5px; color: #334155; line-height: 1.6; padding-left: 11px; position: relative; margin-bottom: 3px; }
.bullets li:before { content: ''; position: absolute; left: 0; top: 6px; width: 4.5px; height: 4.5px; border-radius: 50%; background: #0D9488; }

/* Projects */
.proj { border: 1px solid #E2E8F0; border-radius: 5px; padding: 9px 12px 9px 14px; margin-bottom: 8px; position: relative; overflow: hidden; }
.proj:before { content: ''; position: absolute; left: 0; top: 0; bottom: 0; width: 3px; background: #0D9488; }
.proj-head { display: flex; justify-content: space-between; align-items: center; margin-bottom: 3px; }
.proj-title { font-size: 11.5px; font-weight: 700; }
.proj-store { font-size: 9px; font-weight: 600; color: #0D9488; }
.proj-desc { font-size: 10px; color: #334155; line-height: 1.55; margin-bottom: 4px; }
.proj-stack { font-size: 9.5px; color: #64748B; }
.proj-stack b { color: #0F172A; }
"""

HTML_BODY = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<style>{HTML_CSS}</style>
</head>
<body>

<div class="header">
  <div>
    <div class="h-name">Mohit <span>Kapruwan</span></div>
    <div class="h-sub">Flutter Developer &nbsp;·&nbsp; Mobile &amp; Cross-Platform</div>
  </div>
  <div class="h-contacts">
    <div>mohitkapruwan4@gmail.com &nbsp;·&nbsp; +91 97199 72748</div>
    <div>Mohali, India</div>
    <div>linkedin.com/in/mohit-kapruwan-96b62b22b</div>
    <div>github.com/Mohit-Kapruwan-MK</div>
  </div>
</div>

<div class="body">
  <!-- SIDEBAR -->
  <div class="sidebar">

    <div class="s-section">
      <div class="s-title">Technical Skills</div>
      {"".join(f'''
      <div class="skill">
        <div class="skill-row">
          <span class="skill-name">{n}</span>
          <span class="skill-pct">{p}%</span>
        </div>
        <div class="track"><div class="fill" style="width:{p}%"></div></div>
      </div>
      ''' for n, p in SKILLS)}
    </div>

    <div class="s-section">
      <div class="s-title">Tools &amp; Stack</div>
      <div class="tags">
        {"".join(f'<span class="tag">{t}</span>' for t in ["Flutter","Dart","Firebase","Riverpod","Provider"])}
        {"".join(f'<span class="tag-g">{t}</span>' for t in ["React","Next.js","Node.js","Express","MongoDB","MySQL","TypeScript","SQLite","GraphQL","Git"])}
      </div>
    </div>

    <div class="s-section">
      <div class="s-title">Engineering Practices</div>
      <div class="tags">
        {"".join(f'<span class="tag-g">{p}</span>' for p in PRACTICES)}
      </div>
    </div>

    <div class="s-section">
      <div class="s-title">Languages</div>
      {"".join(f'<div class="lang-row"><span class="lang-name">{l}</span><span class="lang-lv">{lv}</span></div>' for l, lv in LANGUAGES)}
    </div>

    <div class="s-section">
      <div class="s-title">Availability</div>
      <div style="font-size:10.5px;line-height:1.6;color:#0F172A;">
        Open to <b>Freelance</b> &amp; <b>Full-time</b> opportunities worldwide.
      </div>
    </div>

  </div>

  <!-- MAIN -->
  <div class="main">

    <div class="m-section">
      <div class="m-title">Professional Summary</div>
      <div class="summary">{SUMMARY}</div>
    </div>

    <div class="m-section">
      <div class="m-title">Work Experience</div>
      {"".join(f'''
      <div class="exp">
        <div class="exp-head">
          <span class="exp-role">{e["role"]}</span>
          <span class="exp-period">{e["period"]}</span>
        </div>
        <div class="exp-co">{e["company"]} &nbsp;·&nbsp; {e["location"]}</div>
        <ul class="bullets">
          {"".join(f"<li>{b}</li>" for b in e["bullets"])}
        </ul>
      </div>
      ''' for e in EXPERIENCE)}
    </div>

    <div class="m-section">
      <div class="m-title">Key Projects</div>
      {"".join(f'''
      <div class="proj">
        <div class="proj-head">
          <span class="proj-title">{p["title"]}</span>
          <span class="proj-store">&#9654; {p["store"]}</span>
        </div>
        <div class="proj-desc">{p["desc"]}</div>
        <div class="proj-stack"><b>Stack:</b> {p["stack"]}</div>
      </div>
      ''' for p in PROJECTS)}
    </div>

  </div>
</div>

</body>
</html>"""


def build_pdf():
    try:
        from weasyprint import HTML, CSS
        print("Generating PDF...")
        HTML(string=HTML_BODY, base_url=RESUME_DIR).write_pdf(PDF_OUT)
        print(f"  ✓  PDF saved → {PDF_OUT}")
    except Exception as exc:
        print(f"  ✗  PDF failed: {exc}")


# ─────────────────────────────────────────────────────────────────────────────
#  2. DOCX via python-docx
# ─────────────────────────────────────────────────────────────────────────────

def build_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        print("Generating Word document...")

        ACCENT   = RGBColor(0x0D, 0x94, 0x88)
        DARK     = RGBColor(0x0F, 0x17, 0x2A)
        MUTED    = RGBColor(0x64, 0x74, 0x8B)
        BODY_CLR = RGBColor(0x33, 0x41, 0x55)
        WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

        def set_font(run, size, bold=False, color=None, italic=False):
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color

        def add_paragraph(parent, text="", style=None):
            if style:
                p = parent.add_paragraph(style=style)
            else:
                p = parent.add_paragraph()
            if text:
                p.add_run(text)
            return p

        def para_space(p, before=0, after=0):
            p.paragraph_format.space_before = Pt(before)
            p.paragraph_format.space_after  = Pt(after)

        def shade_cell(cell, hex_color):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_color)
            tcPr.append(shd)

        def cell_border(cell, sides, color="0D9488", sz="6"):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in sides:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    sz)
                el.set(qn("w:color"), color)
                tcBorders.append(el)
            tcPr.append(tcBorders)

        def no_border_table(table):
            tbl = table._tbl
            tblPr = tbl.tblPr
            tblBorders = OxmlElement("w:tblBorders")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "none")
                el.set(qn("w:sz"),    "0")
                el.set(qn("w:color"), "auto")
                tblBorders.append(el)
            tblPr.append(tblBorders)

        doc = Document()

        # ── Page margins
        for section in doc.sections:
            section.top_margin    = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin   = Cm(1.6)
            section.right_margin  = Cm(1.6)

        # ── HEADER ───────────────────────────────────────────────
        header_tbl = doc.add_table(rows=1, cols=2)
        no_border_table(header_tbl)
        header_tbl.columns[0].width = Cm(11)

        lc = header_tbl.cell(0, 0)
        rc = header_tbl.cell(0, 1)

        # Name
        np = lc.add_paragraph()
        para_space(np, before=0, after=2)
        nr = np.add_run("Mohit ")
        set_font(nr, 26, bold=True, color=DARK)
        nr2 = np.add_run("Kapruwan")
        set_font(nr2, 26, bold=True, color=ACCENT)

        # Title
        tp = lc.add_paragraph()
        para_space(tp, before=0, after=0)
        tr = tp.add_run("Flutter Developer  ·  Mobile & Cross-Platform")
        set_font(tr, 9.5, color=MUTED)

        # Right column — contact
        rc.paragraphs[0].clear()
        for line in [
            CONTACT["email"],
            CONTACT["phone"] + "   |   " + CONTACT["location"],
            CONTACT["linkedin"],
            CONTACT["github"],
        ]:
            cp = rc.add_paragraph()
            para_space(cp, before=0, after=1)
            cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cr = cp.add_run(line)
            set_font(cr, 9, color=MUTED)

        # Accent rule under header
        rule = doc.add_paragraph()
        para_space(rule, before=4, after=8)
        pPr  = rule._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:color"), "0D9488")
        pBdr.append(bot)
        pPr.append(pBdr)

        # ── BODY (two-column table) ───────────────────────────────
        body_tbl = doc.add_table(rows=1, cols=2)
        no_border_table(body_tbl)
        body_tbl.columns[0].width = Cm(6.2)

        left  = body_tbl.cell(0, 0)
        right = body_tbl.cell(0, 1)
        shade_cell(left, "F8FAFC")
        cell_border(left, ["right"], color="E2E8F0", sz="6")
        left._tc.get_or_add_tcPr().append(
            OxmlElement("w:tcW")
        )

        def section_title_left(cell, text):
            p = cell.add_paragraph()
            para_space(p, before=10, after=4)
            r = p.add_run(text.upper())
            set_font(r, 7.5, bold=True, color=ACCENT)
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "4")
            bot.set(qn("w:color"), "E2E8F0")
            pBdr.append(bot)
            pPr.append(pBdr)

        def section_title_right(cell, text):
            p = cell.add_paragraph()
            para_space(p, before=10, after=4)
            r = p.add_run(text.upper())
            set_font(r, 7.5, bold=True, color=ACCENT)
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "4")
            bot.set(qn("w:color"), "E2E8F0")
            pBdr.append(bot)
            pPr.append(pBdr)

        # ── LEFT: Skills ────────────────────
        section_title_left(left, "Technical Skills")
        for name, pct in SKILLS:
            sp = left.add_paragraph()
            para_space(sp, before=3, after=1)
            sr1 = sp.add_run(name)
            set_font(sr1, 9, bold=True, color=DARK)
            sr2 = sp.add_run(f"  {pct}%")
            set_font(sr2, 8, color=MUTED)

            # Progress bar as a tiny table
            bar_tbl = left.add_table(rows=1, cols=2)
            no_border_table(bar_tbl)
            filled = round(pct / 100 * 10)
            bar_tbl.columns[0].width = Cm(round(pct / 100 * 4.5, 2))
            bar_tbl.columns[1].width = Cm(round((100 - pct) / 100 * 4.5, 2))
            shade_cell(bar_tbl.cell(0, 0), "0D9488")
            shade_cell(bar_tbl.cell(0, 1), "E2E8F0")
            for c in bar_tbl.columns:
                for cell in c.cells:
                    p = cell.paragraphs[0]
                    para_space(p, before=0, after=0)
                    p.paragraph_format.line_spacing = Pt(4)

            gap = left.add_paragraph()
            para_space(gap, before=0, after=3)

        # ── LEFT: Tools ─────────────────────
        section_title_left(left, "Tools & Stack")
        tools_p = left.add_paragraph()
        para_space(tools_p, before=2, after=0)
        set_font(tools_p.add_run(", ".join(TOOLS)), 9, color=BODY_CLR)

        # ── LEFT: Practices ─────────────────
        section_title_left(left, "Engineering Practices")
        for pr in PRACTICES:
            pp = left.add_paragraph()
            para_space(pp, before=1, after=0)
            pr_r = pp.add_run(f"• {pr}")
            set_font(pr_r, 9, color=BODY_CLR)

        # ── LEFT: Languages ──────────────────
        section_title_left(left, "Languages")
        for lang, lv in LANGUAGES:
            lp = left.add_paragraph()
            para_space(lp, before=2, after=0)
            set_font(lp.add_run(lang), 9.5, bold=True, color=DARK)
            set_font(lp.add_run(f"  —  {lv}"), 9, color=MUTED)

        # ── LEFT: Availability ──────────────
        section_title_left(left, "Availability")
        avp = left.add_paragraph()
        para_space(avp, before=2, after=0)
        set_font(avp.add_run("Open to "), 9, color=DARK)
        set_font(avp.add_run("Freelance"), 9, bold=True, color=ACCENT)
        set_font(avp.add_run(" & "), 9, color=DARK)
        set_font(avp.add_run("Full-time"), 9, bold=True, color=ACCENT)
        set_font(avp.add_run(" worldwide."), 9, color=DARK)

        # ── RIGHT: Summary ──────────────────
        section_title_right(right, "Professional Summary")
        sp = right.add_paragraph()
        para_space(sp, before=2, after=0)
        sr = sp.add_run(SUMMARY)
        set_font(sr, 9.5, color=BODY_CLR)
        sp.paragraph_format.line_spacing = Pt(15)

        # ── RIGHT: Experience ───────────────
        section_title_right(right, "Work Experience")
        for exp in EXPERIENCE:
            # Role + period row
            et = right.add_table(rows=1, cols=2)
            no_border_table(et)
            ep_left  = et.cell(0, 0)
            ep_right = et.cell(0, 1)

            rp = ep_left.add_paragraph()
            para_space(rp, before=4, after=0)
            set_font(rp.add_run(exp["role"]), 11, bold=True, color=DARK)

            pp2 = ep_right.add_paragraph()
            para_space(pp2, before=4, after=0)
            pp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_font(pp2.add_run(exp["period"]), 9, bold=True, color=ACCENT)

            # Company
            cp = right.add_paragraph()
            para_space(cp, before=1, after=4)
            set_font(cp.add_run(f"{exp['company']}  ·  {exp['location']}"), 9.5, color=MUTED)

            # Bullets
            for b in exp["bullets"]:
                bp = right.add_paragraph()
                para_space(bp, before=0, after=2)
                bp.paragraph_format.left_indent = Cm(0.3)
                br = bp.add_run(f"• {b}")
                set_font(br, 9, color=BODY_CLR)
                bp.paragraph_format.line_spacing = Pt(13)

            gap = right.add_paragraph()
            para_space(gap, before=0, after=4)

        # ── RIGHT: Projects ─────────────────
        section_title_right(right, "Key Projects")
        for proj in PROJECTS:
            # Project box via table
            pt = right.add_table(rows=1, cols=1)
            no_border_table(pt)
            pc = pt.cell(0, 0)
            cell_border(pc, ["top","right","bottom"], color="E2E8F0", sz="4")
            cell_border(pc, ["left"], color="0D9488", sz="12")

            # Title + store
            ph = pc.add_paragraph()
            para_space(ph, before=2, after=1)
            set_font(ph.add_run(proj["title"]), 10.5, bold=True, color=DARK)
            set_font(ph.add_run(f"   ▸ {proj['store']}"), 8.5, color=ACCENT)

            # Desc
            pd = pc.add_paragraph()
            para_space(pd, before=0, after=2)
            set_font(pd.add_run(proj["desc"]), 9, color=BODY_CLR)
            pd.paragraph_format.line_spacing = Pt(13)

            # Stack
            ps = pc.add_paragraph()
            para_space(ps, before=0, after=4)
            set_font(ps.add_run("Stack: "), 8.5, bold=True, color=DARK)
            set_font(ps.add_run(proj["stack"]), 8.5, color=MUTED)

            gap = right.add_paragraph()
            para_space(gap, before=0, after=2)

        doc.save(DOCX_OUT)
        print(f"  ✓  Word doc saved → {DOCX_OUT}")

    except Exception as exc:
        import traceback
        traceback.print_exc()
        print(f"  ✗  DOCX failed: {exc}")


# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_pdf()
    build_docx()
    print("\nDone! Both files are in assets/resume/")
